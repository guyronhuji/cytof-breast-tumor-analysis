"""Generate CyTOF_Data_Submission_Instructions.docx."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


# ── helpers ────────────────────────────────────────────────────────────────


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_para(doc, text="", style=None, bold=False, italic=False, size=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_code(doc, text):
    """Monospace block styled as 'No Spacing'."""
    for line in text.splitlines():
        p = doc.add_paragraph(style="No Spacing")
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
    return None


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
        # shade header
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "DEEAF1")
        tcPr.append(shd)

    # data rows
    for r, row in enumerate(rows):
        cells = table.rows[r + 1].cells
        for i, cell_text in enumerate(row):
            cells[i].text = str(cell_text)

    # column widths
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Inches(width)

    return table


def add_note_box(doc, text, color="FFF2CC"):
    """Yellow (or custom) shaded paragraph as a note box."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.italic = True
    # shade paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    pPr.append(shd)
    return p


# ── document ───────────────────────────────────────────────────────────────

doc = Document()

# page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Title ──────────────────────────────────────────────────────────────────

title = doc.add_heading("CyTOF Data Submission & Analysis Guide", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(
    "Prepared for lab use — CyTOF standard analysis package (cytofstandard)"
).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# ── Quick Reference Box ────────────────────────────────────────────────────

add_note_box(
    doc,
    "QUICK REFERENCE\n"
    "What to submit for every CyTOF run:\n"
    "  1. Debarcoded data files (FCS, CSV, or Parquet) after preliminary QC by the CyTOF unit.\n"
    "  2. One file named sample_metadata.csv.\n"
    "Required metadata columns: file_name, sample_id, line_id.\n"
    "Every submitted data file must appear exactly once in the metadata and map to one sample and one line.",
    color="D9EAD3",
)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════
# PART 1 — DATA SUBMISSION
# ═══════════════════════════════════════════════════════════════════════════

add_heading(doc, "PART 1 — Data Submission (wet-lab users)", 1)

# ── 1. Purpose ─────────────────────────────────────────────────────────────

add_heading(doc, "1. Purpose", 2)
add_para(
    doc,
    "This document explains what files and metadata must be provided after each CyTOF run, "
    "and how the cytofstandard analysis package processes them. Following these instructions "
    "ensures that every run can be analyzed reproducibly, compared across future runs, and "
    "traced back to the original data without ambiguity.",
)

# ── 2. Why we use the standardized package ─────────────────────────────────

add_heading(doc, "2. Why we use the standardized package", 2)
add_para(
    doc,
    "Until now, CyTOF analysis depended on manually organized files, inconsistent marker names, "
    "and run-specific scripts. The cytofstandard package solves these problems.",
)

add_heading(doc, "2.1  Every cell gets a permanent ID", 3)
add_para(
    doc,
    "During ingestion every cell/event receives a UUID derived from the project, run, source file, "
    "and event index. This ID is fixed forever — if we reanalyze the same run later we can identify "
    "the exact same cells again.",
)

add_heading(doc, "2.2  Marker names are standardized", 3)
add_para(
    doc,
    "CyTOF files often contain slightly different names for the same marker. "
    "The package maps alternate names to the lab's official list. Examples:",
)
add_table(
    doc,
    ["Alternate name", "Standard name"],
    [
        ["H3K27Ac", "H3K27ac"],
        ["H3K9Ac", "H3K9ac"],
        ["H2BK5Ac", "H2BK5ac"],
        ["H4K16Ac", "H4K16ac"],
        ["E-cad", "E-cadherin"],
        ["PanKRT", "Pan-KRT"],
        ["DNA1", "DNA"],
        ["DNA2", "DNA"],
    ],
    col_widths=[2.5, 2.5],
)
doc.add_paragraph()

add_heading(doc, "2.3  Each run is analyzed separately first", 3)
add_para(
    doc,
    "Each CyTOF run is treated as its own analysis unit. Only later, if scientifically appropriate, "
    "runs may be integrated. A run can also be a subset of a CyTOF acquisition or a combination of "
    "files from different acquisitions (requiring careful batch-correction consideration).",
)

add_heading(doc, "2.4  Samples and lines are tracked explicitly", 3)
add_para(
    doc,
    "Each submitted data file must link to a sample, and each sample must link to one biological line. "
    "This is essential for normalization, z-scoring, quality control, and interpretation. "
    "Example: sample_01 → MCF7, sample_02 → T47D, sample_03 → MDA-MB-231.",
)

add_heading(doc, "2.5  The analysis is reproducible and traceable", 3)
add_para(doc, "The package stores:")
for item in [
    "Submitted raw / debarcoded values",
    "Standard marker names and the original → standard mapping",
    "Sample metadata and line information",
    "File hashes and cell UUIDs",
    "Complete processing history with package version",
]:
    add_bullet(doc, item)

# ── 3. Submission workflow ─────────────────────────────────────────────────

add_heading(doc, "3. Submission workflow overview", 2)
add_para(doc, "Follow these steps in order after each CyTOF run:")

steps = [
    ("Step 1", "Receive debarcoded files from the CyTOF unit after preliminary QC."),
    ("Step 2", "Organize files in one folder (one folder per run)."),
    (
        "Step 3",
        "Create sample_metadata.csv — one row per submitted data file with at minimum "
        "file_name, sample_id, and line_id.",
    ),
    (
        "Step 4",
        "Verify: every submitted file appears exactly once in the metadata, "
        "names match exactly (case-sensitive), and required fields are not blank.",
    ),
    ("Step 5", "Share the run folder (data_files/ + sample_metadata.csv) with the analysis person."),
    (
        "Step 6",
        "The analysis person runs the package. If the package reports an error, fix the "
        "metadata or file issue and resubmit. The package stops early rather than silently "
        "analyzing mislabeled data.",
    ),
]

table = doc.add_table(rows=len(steps), cols=2)
table.style = "Table Grid"
for i, (step, desc) in enumerate(steps):
    cells = table.rows[i].cells
    cells[0].text = step
    cells[0].paragraphs[0].runs[0].bold = True
    cells[1].text = desc
    cells[0].width = Inches(0.9)
    cells[1].width = Inches(5.1)
doc.add_paragraph()

# ── 4. Required folder structure ───────────────────────────────────────────

add_heading(doc, "4. Required folder structure", 2)
add_para(doc, "Organize each run in one folder:")
add_code(
    doc,
    """run_2026_05_15/
    data_files/
        sample_A.fcs
        sample_B.fcs
        sample_C.fcs
    sample_metadata.csv""",
)
doc.add_paragraph()
add_para(
    doc,
    "Do not rename data files after preparing sample_metadata.csv. "
    "File names in the metadata must exactly match the submitted file names.",
)

# ── 5. Data file requirements ──────────────────────────────────────────────

add_heading(doc, "5. Data file requirements", 2)

add_heading(doc, "5.1  Accepted file formats", 3)
add_para(doc, "The package accepts three formats:")
add_table(
    doc,
    ["Format", "Extension", "Notes"],
    [
        ["FCS", ".fcs", "Standard flow cytometry format. Preferred when available."],
        ["CSV", ".csv", "One row per cell, one column per marker/channel."],
        ["Parquet", ".parquet", "Columnar format. Treated identically to CSV."],
    ],
    col_widths=[1.2, 1.0, 4.0],
)
doc.add_paragraph()

add_heading(doc, "5.2  Required processing state", 3)
add_para(
    doc,
    "Submit debarcoded files after preliminary QC by the CyTOF unit. "
    "Do not submit the original multiplexed/pre-debarcoding files unless specifically requested.",
)

add_heading(doc, "5.3  CSV / Parquet format", 3)
add_para(
    doc,
    "CSV and Parquet files must have one row per cell/event and one column per marker/channel. "
    "Numeric columns are treated as marker data. Non-numeric columns are preserved but ignored "
    "during analysis. Avoid unnecessary extra columns.",
)
add_para(doc, "Example CSV:")
add_code(
    doc,
    """H3,H3K27ac,H3K9me2,E-cadherin,Pan-KRT,DNA
1.23,0.45,0.87,3.21,2.90,5.20
1.11,0.39,0.91,3.45,3.01,5.18""",
)
doc.add_paragraph()

# ── 6. Sample metadata file ────────────────────────────────────────────────

add_heading(doc, "6. Sample metadata file (sample_metadata.csv)", 2)
add_para(
    doc,
    "Every run must include sample_metadata.csv. This tells the package what each submitted "
    "data file represents. The metadata may contain extra rows for planned, missing, failed, "
    "or excluded samples. Every actual submitted data file must map to exactly one sample "
    "and one line.",
)

add_heading(doc, "6.1  Required columns", 3)
add_table(
    doc,
    ["Column", "Required?", "Meaning"],
    [
        ["file_name", "Yes", "Exact submitted file name. Every submitted file must appear exactly once."],
        ["sample_id", "Yes", "Unique sample identifier for this run."],
        ["line_id", "Yes", "Biological line or group the sample belongs to."],
    ],
    col_widths=[1.4, 1.0, 4.0],
)
doc.add_paragraph()

add_heading(doc, "6.2  Recommended optional columns", 3)
add_table(
    doc,
    ["Column", "Meaning", "Examples"],
    [
        ["condition", "Treatment or experimental condition", "control, treated, vehicle, drug_A"],
        ["replicate_id", "Biological or technical replicate", "rep1, rep2, bio_rep1"],
        ["batch_id", "Wet-lab batch that may affect data", "stain_batch_1, culture_batch_2026_05"],
        ["acquisition_order", "Order on the instrument", "1, 2, 3, 4"],
        ["barcoding_id", "Barcode used for multiplexing", "BC01, BC02"],
        ["notes", "Free-text notes", "low cell number, possible clog"],
    ],
    col_widths=[1.4, 2.8, 2.0],
)
doc.add_paragraph()

add_heading(doc, "6.3  file_name rules", 3)
add_table(
    doc,
    ["Correct", "Incorrect"],
    [
        ["sample_A.fcs", "Sample_A.fcs   (wrong case)"],
        ["sample_A.csv", "sample_A   (missing extension)"],
        ["S001_MCF7_control_rep1.fcs", "/path/to/S001_MCF7_control_rep1.fcs   (full path)"],
    ],
    col_widths=[3.0, 3.0],
)
doc.add_paragraph()

add_heading(doc, "6.4  Example sample_metadata.csv", 3)
add_code(
    doc,
    """file_name,sample_id,line_id,condition,replicate_id,batch_id,acquisition_order,barcoding_id,notes
S001_MCF7_control_rep1.fcs,S001,MCF7,control,rep1,stain_batch_1,1,BC01,
S002_MCF7_treated_rep1.fcs,S002,MCF7,treated,rep1,stain_batch_1,2,BC02,
S003_T47D_control_rep1.fcs,S003,T47D,control,rep1,stain_batch_1,3,BC03,
S004_T47D_treated_rep1.fcs,S004,T47D,treated,rep1,stain_batch_1,4,BC04,""",
)
doc.add_paragraph()

# ── 7. Metadata rules ──────────────────────────────────────────────────────

add_heading(doc, "7. Rules for the metadata file", 2)

rules = [
    (
        "Every submitted data file must appear in the metadata",
        "If data_files/ contains sample_A.fcs, sample_B.fcs, and sample_C.fcs then "
        "sample_metadata.csv must contain rows for all three file names.",
    ),
    (
        "Every submitted file maps to exactly one sample and one line",
        "A file_name that appears twice with different sample_ids or line_ids is an error. "
        "Extra metadata rows for non-submitted samples are fine.",
    ),
    (
        "File names must match exactly",
        "sample_A.fcs, Sample_A.fcs, sample_A.FCS, and sample_A final.fcs are four different names.",
    ),
    (
        "Required fields cannot be blank for submitted files",
        "file_name, sample_id, and line_id must be filled for every submitted file.",
    ),
    (
        "Use consistent spelling",
        "Do not mix MCF7 / MCF-7 / mcf7, or control / Control / CTRL. "
        "Pick one spelling per line or condition and use it everywhere.",
    ),
]

for title_text, desc in rules:
    p = doc.add_paragraph()
    run = p.add_run(title_text + ": ")
    run.bold = True
    p.add_run(desc)

# ── 8. Marker names ────────────────────────────────────────────────────────

add_heading(doc, "8. Marker names", 2)
add_para(
    doc,
    "Use the lab's standard marker names whenever possible. The package recognizes common "
    "alternate names (see section 2.2). If in doubt, use the standard name directly.",
)

add_heading(doc, "8.1  Warning: duplicate DNA channels", 3)
add_note_box(
    doc,
    "If a run has two DNA channels (e.g. DNA1 and DNA2) both may map to the standard name DNA, "
    "which is ambiguous. If this applies to your run, notify the analysis person before ingestion. "
    "We will decide whether to keep them separate, combine them, or rename them.",
    color="FCE5CD",
)
doc.add_paragraph()

# ── 9. Recommended naming conventions ─────────────────────────────────────

add_heading(doc, "9. Recommended naming conventions", 2)

add_heading(doc, "File names", 3)
add_code(
    doc,
    """S001_MCF7_control_rep1.fcs
S002_MCF7_treated_rep1.fcs
S003_T47D_control_rep1.fcs
S004_T47D_treated_rep1.fcs""",
)
doc.add_paragraph()

add_heading(doc, "Sample IDs", 3)
add_code(doc, "S001, S002, S003  —  or  —  MCF7_control_rep1, T47D_treated_rep1")
doc.add_paragraph()

add_heading(doc, "Line IDs", 3)
add_code(doc, "MCF7, T47D, BT474, MDA-MB-231")
doc.add_paragraph()

add_heading(doc, "Conditions", 3)
add_code(doc, "control, treated, vehicle, drug_A, drug_B, day_0")
doc.add_paragraph()

# ── 10. What NOT to do ─────────────────────────────────────────────────────

add_heading(doc, "10. What not to do", 2)
for item in [
    "Do not submit pre-debarcoding multiplexed files unless specifically requested.",
    "Do not manually edit FCS files.",
    "Do not rename data files after creating sample_metadata.csv.",
    "Do not leave blank cells in required metadata columns for submitted files.",
    "Do not use inconsistent spellings for the same line, sample, or condition.",
    "Do not use spaces or unusual special characters in sample IDs if avoidable.",
]:
    add_bullet(doc, item)

# ── 11. Submission checklist ───────────────────────────────────────────────

add_heading(doc, "11. Submission checklist", 2)
add_note_box(
    doc,
    "Use this checklist before sharing your run folder with the analysis person.",
    color="D9EAD3",
)
doc.add_paragraph()

checklist = [
    "Submitted files are debarcoded after preliminary QC by the CyTOF unit.",
    "All data files are in the data_files/ subfolder.",
    "The metadata file is named sample_metadata.csv.",
    "Every submitted data file appears exactly once in sample_metadata.csv.",
    "Each submitted data file maps to one sample_id and one line_id.",
    "file_name entries match submitted file names exactly (case-sensitive).",
    "Every submitted file row has sample_id and line_id filled.",
    "Conditions are spelled consistently throughout.",
    "Replicates are labeled clearly.",
    "acquisition_order is included if known.",
    "Any acquisition problems are described in notes.",
    "No duplicate DNA channels are present, or the analysis person has been informed.",
]
for item in checklist:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("☐  " + item)

# ═══════════════════════════════════════════════════════════════════════════
# PART 2 — PACKAGE OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════

doc.add_page_break()
add_heading(doc, "PART 2 — Package Overview", 1)

# ── 12. What happens after submission ──────────────────────────────────────

add_heading(doc, "12. What happens after you submit the files", 2)
add_para(doc, "The analysis package will:")
for i, step in enumerate([
    "Check that every submitted file is represented in the metadata.",
    "Check that each submitted file maps to exactly one sample and one line.",
    "Read each FCS, CSV, or Parquet file.",
    "Check that marker names match the lab standard list or known alternate names.",
    "Convert accepted alternate marker names to standard names.",
    "Assign a permanent UUID to every cell/event.",
    "Store the run as a standardized AnnData analysis object.",
    "Save the file manifest, marker mapping, sample table, and processing log.",
], 1):
    p = doc.add_paragraph(style="List Number")
    p.add_run(step)

doc.add_paragraph()
add_note_box(
    doc,
    "If something is wrong the package stops and reports the problem. "
    "This is intentional — it is better to stop early than to silently analyze mislabeled data.",
    color="FCE5CD",
)
doc.add_paragraph()

# ── 13. Data storage (AnnData) ─────────────────────────────────────────────

add_heading(doc, "13. How the package stores the data: AnnData / Zarr format", 2)
add_para(
    doc,
    "Each ingested run is stored as one AnnData object in a Zarr store at "
    "runs/<run_id>/processed/<run_id>.zarr. AnnData is the standard format used "
    "in single-cell analysis — it keeps the cell-by-marker matrix together with "
    "cell metadata, marker metadata, embeddings, and processing history in one "
    "self-describing on-disk object.",
)

add_heading(doc, "13.1  Conceptual layout", 3)
add_para(
    doc,
    "Think of AnnData as a spreadsheet where rows are cells and columns are markers, "
    "with additional tables attached for annotation and results:",
)
add_code(
    doc,
    """                      markers (var)
                  H3   H3K27ac  H3K27me3  ECad  ...
               ┌─────────────────────────────────┐
    cell_0001  │  1.23   0.45     0.87    3.21  │  ← obs (cell metadata)
    cell_0002  │  1.11   0.39     0.91    3.45  │    sample_id, line_id,
    cell_0003  │  1.35   0.50     0.76    3.10  │    source_file, cluster ...
    ...        │  ...    ...      ...     ...   │
               └─────────────────────────────────┘
                            ↑ X (active matrix)

    layers["raw"]           — original debarcoded values (never modified)
    layers["normalized"]    — after cytof_transform correction
    layers["normalized_z"]  — after balanced z-scoring

    obsm["X_umap_norm"]     — UMAP coordinates (n_cells × 2)
    obsp["norm_umap_connectivities"] — KNN graph for Leiden clustering

    uns["normalization"]    — normalization parameters and settings
    uns["provenance"]       — full processing history""",
)
doc.add_paragraph()

add_heading(doc, "13.2  obs — cell metadata", 3)
add_para(
    doc,
    "Every row of obs corresponds to one cell. Columns are added progressively "
    "as analysis proceeds:",
)
add_table(
    doc,
    ["obs column", "Added when", "Description"],
    [
        ["cell_uuid", "Ingestion", "Permanent UUID derived from project + run + file + event index."],
        ["project_id", "Ingestion", "Project identifier."],
        ["run_id", "Ingestion", "Run identifier."],
        ["sample_id", "Ingestion", "From sample_metadata.csv."],
        ["line_id", "Ingestion", "From sample_metadata.csv."],
        ["source_file", "Ingestion", "Original filename the cell came from."],
        ["event_index", "Ingestion", "Row index in the original source file (0-based)."],
        ["condition", "Ingestion", "From sample_metadata.csv (if provided)."],
        ["replicate_id", "Ingestion", "From sample_metadata.csv (if provided)."],
        ["batch_id", "Ingestion", "From sample_metadata.csv (if provided)."],
        ["qc_pass", "QC gating", "Boolean: True if the cell passed all QC gates."],
        ["<embedding>_leiden", "Clustering", "Integer cluster label, e.g. norm_umap_leiden."],
        ["PC_<signature>", "PermCell", "Smoothed PermCell z-score for the named signature."],
        ["PC_R_<signature>", "PermCell", "Raw (unsmoothed) PermCell z-score."],
    ],
    col_widths=[1.6, 1.3, 3.3],
)
doc.add_paragraph()

add_heading(doc, "13.3  var — marker metadata", 3)
add_para(
    doc,
    "Every row of var corresponds to one marker. The index is the standard marker name.",
)
add_table(
    doc,
    ["var column", "Description"],
    [
        ["index (var_names)", "Standard marker name (e.g. H3K27ac)."],
        ["original_names", "Comma-separated list of names found in source files before standardization."],
        ["marker_class", "Category from standard_markers.csv (e.g. histone_core, histone_mark, identity)."],
        ["intra", "True if marker is classified as intra-nuclear."],
        ["extra", "True if marker is classified as extra-nuclear / surface."],
    ],
    col_widths=[1.8, 4.4],
)
doc.add_paragraph()

add_heading(doc, "13.4  X and layers — expression matrices", 3)
add_para(
    doc,
    "X is the active cell × marker expression matrix. Layers store named snapshots "
    "of the matrix at different processing stages. The original raw data is always "
    "preserved in layers['raw'] and is never overwritten.",
)
add_table(
    doc,
    ["Layer", "Contents", "Created by"],
    [
        ["raw", "Original debarcoded values as submitted. Never modified after ingestion.", "run.ingest()"],
        ["normalized", "Corrected marker values in arcsinh space after cytof_transform.", "run.normalize_with_cytof_transform()"],
        ["normalized_z", "Balanced z-scored values computed from the normalized layer.", "run.zscore_markers_balanced()"],
        ["X (active)", "Points to whichever layer is selected as the default for analysis.", "run.set_x_from_layer()"],
    ],
    col_widths=[1.5, 3.3, 2.0],
)
doc.add_paragraph()
add_note_box(
    doc,
    "X is a pointer, not a separate copy. Calling run.set_x_from_layer('normalized') makes X "
    "point to the normalized values without duplicating storage. The raw layer is always available "
    "as layers['raw'] regardless of what X points to.",
    color="D9EAD3",
)
doc.add_paragraph()

add_heading(doc, "13.5  obsm — per-cell multidimensional results", 3)
add_para(
    doc,
    "obsm stores arrays with one row per cell and multiple columns — typically "
    "embedding coordinates. All obsm keys from a named UMAP use the embedding_name "
    "as a prefix.",
)
add_table(
    doc,
    ["obsm key", "Shape", "Contents"],
    [
        ["X_<embedding_name>", "n_cells × 2", "UMAP coordinates (or n_components if changed)."],
        ["<embedding_name>_indices", "n_cells × k", "KNN indices used to build the Leiden graph."],
        ["<embedding_name>_distances", "n_cells × k", "KNN distances."],
        ["<result_prefix>_smoothed", "n_cells × n_signatures", "PermCell smoothed score matrix."],
    ],
    col_widths=[2.0, 1.3, 2.9],
)
doc.add_paragraph()

add_heading(doc, "13.6  obsp — sparse cell × cell matrices", 3)
add_para(
    doc,
    "obsp stores sparse n_cells × n_cells matrices, primarily the KNN connectivity "
    "graphs used by the Leiden clustering step.",
)
add_table(
    doc,
    ["obsp key", "Contents"],
    [
        ["<embedding_name>_connectivities", "KNN adjacency matrix (binary or weighted). Used by cluster_leiden()."],
        ["<embedding_name>_jaccard_connectivities", "Jaccard-weighted graph. Used by cluster_leiden_jaccard()."],
    ],
    col_widths=[2.5, 3.7],
)
doc.add_paragraph()

add_heading(doc, "13.7  uns — run-level metadata and provenance", 3)
add_para(
    doc,
    "uns is a dictionary of unstructured metadata. Everything written by the package "
    "is stored under a descriptive key so it can be inspected later.",
)
add_table(
    doc,
    ["uns key", "Contents"],
    [
        ["cytofstandard_version", "Package version at ingestion time."],
        ["project_id / run_id", "Project and run identifiers."],
        ["marker_mapping", "Dict mapping each standard marker name to the original names found in source files."],
        ["ingestion_settings", "drop_columns, strict_markers, copy_raw, common_markers_only flags used at ingestion."],
        ["normalization", "Every setting of the normalization call: method (regress or divide), control markers, markers corrected, groupby column, cofactor, gamma mode, protected covariates, compartments, and output layer names."],
        ["zscore_<layer>", "Balancing parameters and mean/std values used for z-scoring."],
        ["clusterings", "One entry per cluster_leiden / cluster_leiden_jaccard call with resolution and seed."],
        ["embeddings", "One entry per compute_umap call with marker list, neighbors, and UMAP parameters."],
        ["permcell_<prefix>", "Signatures, positions key, and PermCell settings for each run_permcell call."],
        ["provenance", "JSONL-style log of all processing steps with timestamps and package version."],
    ],
    col_widths=[2.0, 4.2],
)
doc.add_paragraph()

add_heading(doc, "13.8  File layout on disk", 3)
add_para(doc, "Each project is a directory tree:")
add_code(
    doc,
    """my_project/
├── project.json                  — project metadata and run registry
├── standard_markers.csv          — copy of the marker standard file
├── marker_aliases.yaml           — copy of the alias file (if provided)
└── runs/
    └── run_001/
        ├── run.json              — run metadata (id, name, panel, dates, status)
        ├── provenance.jsonl      — chronological log of all processing steps
        ├── raw/                  — original submitted files (if copy_raw=True)
        │   ├── sample_A.fcs
        │   └── sample_B.fcs
        └── processed/
            └── run_001.zarr/     — AnnData stored as Zarr
                ├── X/            — active expression matrix
                ├── obs/          — cell metadata columns
                ├── var/          — marker metadata columns
                ├── layers/
                │   ├── raw/
                │   ├── normalized/
                │   └── normalized_z/
                ├── obsm/         — embeddings and PermCell matrices
                ├── obsp/         — KNN and Jaccard connectivity matrices
                └── uns/          — run metadata and provenance""",
)
doc.add_paragraph()

add_heading(doc, "13.9  Inspecting the AnnData object in Python", 3)
add_code(
    doc,
    """adata = run.read_adata()

print(adata)
# AnnData object with n_obs × n_vars = 45231 × 18
#   obs: 'cell_uuid', 'sample_id', 'line_id', 'source_file', ...
#   var: 'marker_class', 'original_names', 'intra', 'extra'
#   layers: 'raw', 'normalized', 'normalized_z'
#   obsm: 'X_norm_umap', 'norm_umap_indices', 'norm_umap_distances'
#   obsp: 'norm_umap_connectivities'
#   uns: 'cytofstandard_version', 'marker_mapping', 'normalization', ...

# Access specific parts
adata.obs["sample_id"].value_counts()     # cells per sample
adata.var_names                           # standard marker names
adata.layers["raw"][:5, :3]              # raw values, first 5 cells × 3 markers
adata.obsm["X_norm_umap"][:5]            # UMAP coordinates of first 5 cells
adata.uns["normalization"]["control_markers"]""",
)
doc.add_paragraph()
add_para(
    doc,
    "Wet-lab users do not need to work with AnnData directly. Accurate sample metadata "
    "and consistent marker names are all that is needed to build this object correctly.",
    italic=True,
)

# ═══════════════════════════════════════════════════════════════════════════
# PART 3 — API REFERENCE
# ═══════════════════════════════════════════════════════════════════════════

doc.add_page_break()
add_heading(doc, "PART 3 — API Reference (analysis users)", 1)
add_para(
    doc,
    "This section documents the full cytofstandard API. Follow steps 14–28 in order for "
    "a complete end-to-end analysis.",
)

# ── 14. Installation ───────────────────────────────────────────────────────

add_heading(doc, "14. Installation", 2)
add_code(
    doc,
    "pip install anndata zarr numpy pandas pyyaml fcsparser matplotlib seaborn pyarrow",
)
doc.add_paragraph()
add_para(doc, "Install the package itself (if not on PyPI, install from source):")
add_code(doc, "pip install -e /path/to/CyTOFSTD")
doc.add_paragraph()

# ── 15. Prepare marker registry files ─────────────────────────────────────

add_heading(doc, "15. Prepare marker registry files", 2)
add_para(doc, "You need two files:")
add_bullet(doc, "standard_markers.csv — columns: standard_marker_name, marker_class")
add_bullet(doc, "marker_aliases.yaml — (optional) maps alternate names to standard names")
add_para(doc, "Example standard_markers.csv:")
add_code(
    doc,
    """standard_marker_name,marker_class
H3,histone_core
H3K27ac,histone_mark
H3K27me3,histone_mark
E-cadherin,identity
EpCAM,identity
Vimentin,identity""",
)
doc.add_paragraph()

# ── 16. Create a project ───────────────────────────────────────────────────

add_heading(doc, "16. Create a project", 2)
add_code(
    doc,
    """from cytofstandard import Project

project = Project.create(
    path="my_project",
    project_id="BRCA_CYTOF_2026",
    project_name="BRCA CyTOF histone panel",
    standard_marker_file="standard_markers.csv",
    marker_alias_file="marker_aliases.yaml",   # optional
)""",
)
doc.add_paragraph()
add_para(doc, "To load an existing project:")
add_code(
    doc,
    """project = Project.load("my_project")""",
)
doc.add_paragraph()

add_table(
    doc,
    ["Parameter", "Required?", "Description"],
    [
        ["path", "Yes", "Directory to create / load the project in."],
        ["project_id", "Yes", "Short identifier (e.g. BRCA_CYTOF_2026)."],
        ["project_name", "Yes", "Human-readable name."],
        ["standard_marker_file", "Yes", "Path to standard_markers.csv or .parquet."],
        ["marker_alias_file", "No", "Path to marker_aliases.yaml."],
        ["overwrite", "No", "Set True to recreate an existing project (default False)."],
    ],
    col_widths=[1.6, 0.9, 3.8],
)
doc.add_paragraph()

# ── 17. List and manage runs ───────────────────────────────────────────────

add_heading(doc, "17. List and manage runs", 2)
add_code(
    doc,
    """# List all runs in the project
project.list_runs()          # returns a DataFrame

# Check if a run exists
project.has_run("run_001")

# Get an existing run object
run = project.get_run("run_001")

# Rename a run
project.rename_run("run_001", "First BRCA run (renamed)")

# Delete a run and all its files (irreversible)
project.remove_run("run_001")""",
)
doc.add_paragraph()

# ── 18. Register and ingest a run ─────────────────────────────────────────

add_heading(doc, "18. Register a run and ingest data", 2)
add_code(
    doc,
    """run = project.add_run(
    run_id="run_001",
    run_name="First BRCA CyTOF run",
    panel_id="breast_histone_panel_v1",
    acquisition_date="2026-05-15",
    instrument="Helios",
    operator="GR",
)

run.ingest(
    files=["data_files/sample_A.fcs", "data_files/sample_B.fcs"],
    sample_metadata="sample_metadata.csv",
    copy_raw=True,           # stores original files under runs/<run_id>/raw
    strict_markers=True,     # fail on unknown markers
    allow_extra_markers=False,
    common_markers_only=False,  # set True to keep only markers present in ALL files
    drop_columns=["Time"],   # remove raw channels before marker processing
)""",
)
doc.add_paragraph()

add_table(
    doc,
    ["Parameter", "Default", "Description"],
    [
        ["files", "—", "List of paths to FCS, CSV, or Parquet files."],
        ["sample_metadata", "—", "Path to sample_metadata.csv."],
        ["copy_raw", "True", "Copy original files to runs/<run_id>/raw for archiving."],
        ["strict_markers", "True", "Fail if unknown marker names are found."],
        ["allow_extra_markers", "False", "Allow markers not in the standard list (useful while testing)."],
        ["common_markers_only", "False", "Drop markers absent from any file in this run (useful for inconsistent panels)."],
        ["drop_columns", "None", "Raw channel names to remove before processing (e.g. ['Time', 'Event_length'])."],
    ],
    col_widths=[1.7, 0.8, 3.8],
)
doc.add_paragraph()

# ── 19. Load data ──────────────────────────────────────────────────────────

add_heading(doc, "19. Load and save data", 2)
add_code(
    doc,
    """# Load the AnnData object for this run
adata = run.read_adata()

# After external modifications, persist back to disk
run.save(adata)

# Check ingestion status
run.is_ingested()    # True / False
run.status           # "registered" | "ingested" | "failed_ingestion"

# Path to the Zarr store
run.zarr_path""",
)
doc.add_paragraph()

# ── 20. QC ─────────────────────────────────────────────────────────────────

add_heading(doc, "20. Quality control", 2)
add_code(
    doc,
    """# Histograms per sample for selected markers
run.plot_marker_histograms(
    markers=["H3", "ECad"],
    layer="raw",          # or "X" for the default layer
    cofactor=5.0,         # arcsinh cofactor for display
)

# Apply QC gates — keep cells within bounds
run.qc_gate(
    gates={
        "H3":   {"lower": 100, "upper": 400},   # absolute bounds
        "DNA":  ("p1", "p99"),                   # percentile bounds
    },
    layer="raw",
    inplace=True,         # persist filtered AnnData to disk
)""",
)
doc.add_paragraph()

# ── 21. Normalization ──────────────────────────────────────────────────────

add_heading(doc, "21. Normalization (cytof_transform)", 2)
add_para(
    doc,
    "Normalize histone mark intensities using control markers to remove technical variation. "
    "Requires the external cytof_transform package.",
)
add_code(
    doc,
    """summary = run.normalize_with_cytof_transform(
    control_markers=["H3.3", "H3", "H4"],       # core histones used as technical reference
    markers_to_correct=["H3K27ac", "H3K4me3"],  # marks to normalize
    source_layer="raw",
    corrected_layer="normalized",               # output: arcsinh-space corrected values
    z_layer="normalized_z",                     # output: z-scored corrected values
    groupby_col="sample_id",                    # normalize per sample (or per line_id)
    input_is_arcsinh=False,
    arcsinh_cofactor=5.0,
)""",
)
doc.add_paragraph()
add_para(doc, "QC plots for normalization:")
add_code(
    doc,
    """run.plot_normalization_tech_factor_qc(
    control_markers=["H3.3", "H3", "H4"],
    layer="raw",
)
run.plot_normalization_gamma_qc(group_value="S001")
run.plot_normalization_marker_correlations_qc(
    pre_layer="raw",
    post_layer="normalized",
)""",
)
doc.add_paragraph()

# ── 22. Z-score ────────────────────────────────────────────────────────────

add_heading(doc, "22. Balanced z-score normalization", 2)
add_para(
    doc,
    "Z-score all markers using a balanced subsample across groups to avoid bias "
    "from unequal group sizes.",
)
add_code(
    doc,
    """run.zscore_markers_balanced(
    source_layer="normalized",
    output_layer="normalized_z",
    groupby_col="sample_id",   # balance by sample (or line_id)
    random_state=0,
    inplace=True,
)""",
)
doc.add_paragraph()

# ── 23. Subsetting ─────────────────────────────────────────────────────────

add_heading(doc, "23. Subsets and subsampling", 2)
add_code(
    doc,
    """# Create a persistent subset run from specific samples or lines
subset_run = run.create_subset_run(
    new_run_id="run_001_MCF7",
    sample_ids=["S001", "S002"],    # keep these samples
    line_ids=None,                  # or filter by line_id
    run_name="MCF7 subset of run_001",
    notes="Created for MCF7-only analysis",
)

# In-memory balanced subsample (not persisted as a new run)
adata_sub = run.subsample_by_group(
    groupby_col="sample_id",
    n_per_group=1000,    # None = use the smallest group size
    random_state=0,
)""",
)
doc.add_paragraph()

# ── 24. UMAP / clustering ──────────────────────────────────────────────────

add_heading(doc, "24. Embedding and clustering", 2)
add_code(
    doc,
    """# Set X to the layer you want to embed
run.set_x_from_layer("normalized")

# Standard UMAP (all cells)
run.compute_umap(
    markers=["H3", "H3K27me3", "ECad", "EpCAM"],
    source_layer="normalized",
    embedding_name="norm_umap",
    n_neighbors=15,
    min_dist=0.1,
)

# Balanced UMAP — fit on an equal-size subsample per group, transform all cells
run.compute_umap_balanced(
    markers=["H3", "H3K27me3", "ECad", "EpCAM"],
    source_layer="normalized",
    embedding_name="norm_umap_balanced",
    groupby_col="sample_id",
    n_per_group=2000,     # None = use the smallest group size
)

# Leiden clustering on the KNN graph from compute_umap / compute_umap_balanced
run.cluster_leiden(
    embedding_name="norm_umap",
    resolution=1.0,
)

# PhenoGraph-style: Jaccard similarity + Leiden
run.cluster_leiden_jaccard(
    embedding_name="norm_umap",
    resolution=1.0,
    min_jaccard=0.0,      # prune edges below this Jaccard threshold
)""",
)
doc.add_paragraph()

# ── 25. PermCell ───────────────────────────────────────────────────────────

add_heading(doc, "25. PermCell signature scoring", 2)
add_code(
    doc,
    """import PermCell_Smooth as PCS

signatures = {
    "EPI_like":  {"up": ["EpCAM", "KRT8-18"], "down": ["Vimentin"]},
    "Stem_like": ["CD44", "BMI1"],
}

run.run_permcell(
    signatures=signatures,
    source_layer="raw",
    positions_key="X_umap",
    result_prefix="pc",
    permcell_module=PCS,
)

# Results are written to obs:
#   PC_<signature>   — smoothed z-score
#   PC_R_<signature> — raw (unsmoothed) z-score

# Build an AnnData view of PermCell results for plotting/export
adata_pc = run.permcell_to_adata(result_prefix="pc", score="z", smoothed=True)""",
)
doc.add_paragraph()

# ── 26. Plotting ───────────────────────────────────────────────────────────

add_heading(doc, "26. Plotting", 2)

add_heading(doc, "Heatmap", 3)
add_code(
    doc,
    """run.plot_heatmap(
    fields=["H3", "H3K27me3", "PC_EPI_like"],
    groupby="sample_id",             # or ["line_id", "condition"]
    layer="normalized",
    agg="mean",                      # "mean" or "median"
    standard_scale="row",            # z-score over rows (fields)
)""",
)
doc.add_paragraph()

add_heading(doc, "Boxplot with significance", 3)
add_code(
    doc,
    """run.plot_boxplot(
    field="PC_EPI_like",             # marker or obs column
    groupby="condition",
    layer="X",
    order=["control", "treated"],
    comparisons="all",               # or "adjacent" or list of (A, B) pairs
    test="mannwhitney",              # "mannwhitney", "ttest", or "welch"
    multitest="bh",                  # "bh" (Benjamini-Hochberg) or "bonferroni"
    show_points=True,
    max_points=2000,
)""",
)
doc.add_paragraph()

# ── 27. Statistical comparisons ────────────────────────────────────────────

add_heading(doc, "27. Statistical comparisons", 2)
add_code(
    doc,
    """# Pairwise group comparisons (returns a DataFrame)
table = run.compare_groups(
    field="PC_EPI_like",
    groupby="condition",
    method="ttest",           # "ttest" or "wald"
    equal_var=True,           # Welch correction when False
    comparisons="all",        # "all", "adjacent", or explicit pair list
    multitest="bh",           # "bh" or "bonferroni"
)
# Returns: group_a, group_b, n_a, n_b, mean_a, mean_b, stat, p_value, p_adj, method""",
)
doc.add_paragraph()

# ── 28. Export ─────────────────────────────────────────────────────────────

add_heading(doc, "28. Export for downstream analysis", 2)
add_code(
    doc,
    """# Export selected markers and obs columns to a DataFrame
df = run.to_dataframe(
    fields=["H3", "H3K27me3", "sample_id", "condition"],
    layer="raw",
)
df.to_csv("my_export.csv")""",
)
doc.add_paragraph()

# ── 29. Locking and reproducibility ───────────────────────────────────────

add_heading(doc, "29. Locking raw data for reproducibility", 2)
add_para(
    doc,
    "Lock raw layers to read-only to prevent accidental modification after ingestion.",
)
add_code(
    doc,
    """# Lock specific Zarr paths
run.lock_zarr_parts(parts=["layers/raw", "obs"])

# Lock the entire Zarr store
run.lock_zarr_parts()

# Unlock when updates are needed
run.unlock_zarr_parts(parts=["layers/raw", "obs"])""",
)
doc.add_paragraph()

# ── 30. Common issues ──────────────────────────────────────────────────────

add_heading(doc, "30. Common issues and solutions", 2)
add_table(
    doc,
    ["Issue", "Solution"],
    [
        [
            "MetadataValidationError: file not in metadata",
            "Check that the file_name in sample_metadata.csv matches exactly (case-sensitive, with extension).",
        ],
        [
            "MarkerValidationError: unknown markers",
            "Set allow_extra_markers=True to allow them, or add them to marker_aliases.yaml.",
        ],
        [
            "Missing markers across files",
            "Set common_markers_only=True to keep only markers present in all files in the run.",
        ],
        [
            "Sample metadata mismatch",
            "Ensure every input file appears in sample_metadata and that sample_id maps to exactly one line_id.",
        ],
        [
            "Dense boxplots / slow plotting",
            "Use show_points=False and consider subsetting with create_subset_run() before exploratory plotting.",
        ],
        [
            "UMAP biased by unequal group sizes",
            "Use compute_umap_balanced() with n_per_group set to a reasonable number.",
        ],
    ],
    col_widths=[2.2, 4.0],
)
doc.add_paragraph()

# ── Summary ────────────────────────────────────────────────────────────────

doc.add_page_break()
add_heading(doc, "Summary", 1)
add_para(doc, "For every CyTOF run, provide:")
for item in [
    "Debarcoded FCS, CSV, or Parquet files after preliminary QC by the CyTOF unit.",
    "sample_metadata.csv with at minimum: file_name, sample_id, line_id.",
    "Recommended: condition, replicate_id, batch_id, acquisition_order, barcoding_id, notes.",
]:
    add_bullet(doc, item)

doc.add_paragraph()
add_para(
    doc,
    "The cytofstandard package will ingest the data, standardize marker names, assign "
    "permanent cell UUIDs, and store the run as a reproducible AnnData object ready for "
    "normalization, z-scoring, embedding, clustering, and statistical analysis.",
)
doc.add_paragraph()
add_para(
    doc,
    "Prepared for lab use — cytofstandard package — 2026",
    italic=True,
)

# ── Save ───────────────────────────────────────────────────────────────────

out_path = "CyTOF_Data_Submission_Instructions.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
